from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# --- TITLE ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Brain-Based Mental Health Community')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Website Development — Group Feedback')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Please review and share your thoughts on the following questions.\nYour input will shape the direction of our combined website.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# --- SECTION 1: TITLE ---
p = doc.add_paragraph()
run = p.add_run('1. Website Title')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

p = doc.add_paragraph()
run = p.add_run('We are selecting a title for the combined website. The title should communicate that this resource is for everyone — clinicians, caregivers, and self-advocates — and that understanding the brain is the starting point for effective mental health support.')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Please rank your preference (1 = top choice, 4 = least preferred) or note any reactions:')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()  # spacer

# Title options
options = [
    ('Option A', 'Brain-Based Mental Health: Better Outcomes Start with Understanding the Brain'),
    ('Option B', 'When We Start with the Brain, Everyone Benefits'),
    ('Option C', 'The Brain-Based Approach: Connecting Clinicians, Caregivers, and Self-Advocates'),
    ('Option D', 'It Starts with the Brain: Mental Health That Actually Fits'),
]

for label, title in options:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}:  ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
    run = p.add_run(f'"{title}"')
    run.font.size = Pt(12)
    run.italic = True

    # Rank line
    p = doc.add_paragraph()
    run = p.add_run('    Rank: ___    Comments: ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.add_run('_' * 60).font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    doc.add_paragraph()  # spacer

# Other suggestion
p = doc.add_paragraph()
run = p.add_run('Other title suggestion: ')
run.bold = True
run.font.size = Pt(11)
p.add_run('_' * 70).font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

doc.add_paragraph()
doc.add_paragraph()

# --- SECTION 2: ADDITIONAL FEEDBACK ---
p = doc.add_paragraph()
run = p.add_run('2. Additional Questions')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

p = doc.add_paragraph()
run = p.add_run('(Space reserved for additional questions to be added)')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# Add several blank lines for future questions
for i in range(8):
    doc.add_paragraph()

# --- FOOTER ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Thank you for your feedback!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Please return your responses to: ___________________')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(r'C:\Users\Tamra\brain-based-community\Group-Feedback-Website-Title.docx')
print('Word document created successfully.')
