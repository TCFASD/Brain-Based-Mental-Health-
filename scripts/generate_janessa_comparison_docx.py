"""Generate a Word doc comparing the current website text to Janessa's proposed changes.

For each of her suggestions, the doc shows:
  - CURRENT site text (what's on the page today)
  - JANESSA'S suggestion (her proposed change)
  - STATUS (Done / Partially done / Not yet done)
  - Notes
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "Group Feedback" / "Janessa-Feedback-vs-Website.docx"

GREEN = RGBColor(0x4A, 0x6B, 0x5C)
DARK_GREEN = RGBColor(0x2D, 0x4F, 0x3F)
ORANGE = RGBColor(0xC2, 0x6B, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x2E, 0x2E, 0x2E)
DONE = RGBColor(0x2E, 0x7D, 0x32)
PARTIAL = RGBColor(0xE6, 0x8A, 0x00)
TODO = RGBColor(0xC1, 0x37, 0x2A)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_label_run(p, text, color):
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    return r


def add_status_cell(cell, status):
    """status: done | partial | todo"""
    cell.text = ""
    p = cell.paragraphs[0]
    if status == "done":
        add_label_run(p, "DONE", DONE)
        set_cell_shading(cell, "E8F5E9")
    elif status == "partial":
        add_label_run(p, "PARTIAL", PARTIAL)
        set_cell_shading(cell, "FFF8E1")
    else:
        add_label_run(p, "TO DO", TODO)
        set_cell_shading(cell, "FDECEA")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = GREEN
    if level == 1:
        run.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(15)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if italic:
        r.italic = True
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_change_block(doc, title, current_text, janessa_text, status, notes=None, current_label="CURRENT WEBSITE", suggestion_label="JANESSA'S SUGGESTION"):
    """Render a single change comparison as a 2-column table."""
    # Sub-heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = DARK_GREEN

    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(2.85)
    table.columns[1].width = Inches(2.85)
    table.columns[2].width = Inches(0.95)

    # Header row
    hdr = table.rows[0]
    for cell, label, fill in (
        (hdr.cells[0], current_label, "EFEFEF"),
        (hdr.cells[1], suggestion_label, "FFF4E5"),
        (hdr.cells[2], "STATUS", "EFEFEF"),
    ):
        cell.text = ""
        cp = cell.paragraphs[0]
        cr = cp.add_run(label)
        cr.bold = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = GRAY
        set_cell_shading(cell, fill)

    # Body row
    body = table.rows[1]
    body.cells[0].text = ""
    p0 = body.cells[0].paragraphs[0]
    r0 = p0.add_run(current_text)
    r0.font.size = Pt(10)

    body.cells[1].text = ""
    p1 = body.cells[1].paragraphs[0]
    r1 = p1.add_run(janessa_text)
    r1.font.size = Pt(10)

    add_status_cell(body.cells[2], status)
    # Center status cell vertically/horizontally
    body.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), "CCCCCC")
                borders.append(b)
            tc_pr.append(borders)

    if notes:
        np = doc.add_paragraph()
        np.paragraph_format.space_before = Pt(2)
        np.paragraph_format.space_after = Pt(8)
        nr = np.add_run("Notes: ")
        nr.bold = True
        nr.italic = True
        nr.font.size = Pt(9.5)
        nr.font.color.rgb = GRAY
        nr2 = np.add_run(notes)
        nr2.italic = True
        nr2.font.size = Pt(9.5)
        nr2.font.color.rgb = GRAY


def add_addition_block(doc, title, current_text, addition_text, status, notes=None):
    """For Janessa's 'ADD' suggestions where there's no direct current text to swap."""
    add_change_block(
        doc, title, current_text, addition_text, status, notes,
        current_label="CURRENT WEBSITE",
        suggestion_label="JANESSA'S ADDITION / SUGGESTION",
    )


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    add_heading(doc, "Janessa's Feedback vs. Current Website", 1)
    add_para(
        doc,
        "Side-by-side comparison of Janessa Nikols' suggested changes against the live site at "
        "tcfasd.github.io/Brain-Based-Mental-Health-/. Status: DONE = already on the live site; "
        "PARTIAL = some elements present but not all; TO DO = not yet incorporated.",
        italic=True, color=GRAY, size=10,
    )

    # ========== HOME PAGE ==========
    add_heading(doc, "Home Page (index.html)", 2)

    add_heading(doc, "Title Preference", 3)
    add_change_block(
        doc,
        "Site title (H1)",
        "Brain-Based Mental Health: Better Outcomes Start with Understanding the Brain",
        "TOP CHOICE: Option D — \"It Starts with the Brain: Mental Health That Actually Fits\" "
        "(\"Feels most human and names the problem\"). Current Option A she rated \"Not preferred — "
        "Too textbooky and doesn't scream 'who is this for' and 'why should I care'.\"",
        "todo",
        notes="Decision needed: keep current title (which AK ranked top choice) or switch to Janessa's preferred Option D.",
    )

    add_heading(doc, "Overview Paragraph", 3)
    add_change_block(
        doc,
        "Overview paragraph",
        "This site was built by therapists, caregivers, and self-advocates who saw the same problem "
        "from different sides... [continues with Option A wording]",
        "TOP CHOICE: Option B — \"We are therapists, caregivers, and self-advocates who came together "
        "around a shared realization: the mental health system was not working for the people who needed "
        "it most...\" Loves the lines \"Told to try harder...\" and \"When we start with...\". "
        "Tied with Option C: \"Therapists felt unequipped. Families felt unheard. Self-advocates felt unseen.\" "
        "Suggests removing the last sentence (\"This site gives...\") from Option B.",
        "todo",
    )

    add_heading(doc, "Suggested Alternative Headline", 3)
    add_addition_block(
        doc,
        "Hero headline (subtitle)",
        "Resources, training, and guidance for clinicians, caregivers, and self-advocates "
        "navigating FASD and other neurodevelopmental differences.",
        "\"Mental health care often fails brains that work differently.\" + subtitle: "
        "\"Resources for clinicians, caregivers, and self-advocates navigating FASD, autism, trauma, and more. "
        "Built by therapists, caregivers, and self-advocates who saw mental health systems ignoring how brains actually function.\"",
        "todo",
        notes="Janessa called the current headline \"textbooky\" — wants to lead with the problem.",
    )

    add_heading(doc, "Section Order", 3)
    add_addition_block(
        doc,
        "Page section order",
        "Current order: Hero → Overview → Understanding Brain-Based Differences → Audience Cards "
        "(Clinicians/Caregivers/Self-Advocates) → Why Brain-Based Matters",
        "Move audience cards (Clinicians/Caregivers/Self-Advocates) to the TOP so visitors can take "
        "action quickly. Suggested order: Headline → Audience → Why Brain-Based Matters → "
        "Brain-Based Differences → Credibility → Language Matters. Move \"Why Brain-Based Matters\" UP.",
        "todo",
        notes="Major restructuring of home page flow.",
    )

    add_addition_block(
        doc,
        "Audience cards — remove \"Here you will find\"",
        "Each audience card ends with: \"Here you will find a framework for understanding...\" / "
        "\"Here you will find guidance to understand...\" / \"Here you will find tools to understand...\"",
        "Remove the \"Here you will find\" framing. Lead with the audience's problem instead.",
        "todo",
    )

    add_heading(doc, "Audience Card Tone Tweaks", 3)
    add_addition_block(
        doc,
        "Clinicians one-liner (above existing paragraph)",
        "(No one-line descriptor; current opens with \"Your training gave you strong, evidence-based tools. "
        "But it likely did not teach you to assess or understand the brain first...\")",
        "Add ABOVE the longer text: \"Most training doesn't fully address how to assess brain function first.\" "
        "OR \"Tried everything with complex clients? Assess brain function first.\" "
        "Janessa's note: current tone \"could be defensive.\"",
        "todo",
    )

    add_addition_block(
        doc,
        "Caregivers one-liner",
        "(No one-line descriptor.)",
        "\"Understand your child and demand approaches that actually fit.\" "
        "OR \"Traditional strategies fail your child. Get brain-based guidance.\"",
        "todo",
    )

    add_addition_block(
        doc,
        "Self-Advocates one-liner",
        "(No one-line descriptor; current paragraph: \"You have been told to try therapy, use coping skills...\")",
        "\"Therapy expects 'progress' your brain can't deliver that way. Get tools to name what you need.\" "
        "OR \"Therapy assumes typical brains. Learn yours + speak up.\" "
        "Janessa's note: current is \"clinically sounding.\"",
        "todo",
    )

    add_heading(doc, "Brain-Based Differences Section", 3)
    add_change_block(
        doc,
        "\"Understanding Brain-Based Differences\" section",
        "Full section with paragraph + three info cards (Developmental / Acquired / Trauma-Related). "
        "Quite detailed — includes lists of specific conditions under each card.",
        "Shorten this section. Add a \"Learn More\" link. \"Helpful for clinicians, but doesn't really say "
        "'get help now'.\" Suggested text: \"Brain-based differences affect how people learn, regulate, "
        "process information, and respond to stress. They're not behavioral choices — they require support, "
        "not willpower. Examples include FASD, autism, ADHD, brain injury, and trauma. [Learn More About "
        "Brain-Based Differences]\" — separate page or expandable.",
        "todo",
    )

    add_heading(doc, "Quotes Section (Recommended by Workgroup)", 3)
    add_addition_block(
        doc,
        "Testimonial quotes",
        "(Currently absent.)",
        "Add three quotes: \"I love Janessa\" — Clinician [placeholder]; "
        "\"Less reactiveness and more strategies that help\" — Caregiver; "
        "\"I could explain my brain to providers.\" — Self-advocate",
        "todo",
        notes="Workgroup-recommended additions Janessa flagged.",
    )

    add_heading(doc, "Credibility Statement", 3)
    add_addition_block(
        doc,
        "Credibility line",
        "(Currently absent — credibility implied via the About Us page.)",
        "\"Created by licensed clinicians and neurobehavioral specialists.\" "
        "Janessa: \"Missing and important for professionals.\"",
        "todo",
    )

    add_heading(doc, "Language Matters Section", 3)
    add_addition_block(
        doc,
        "Language Matters section + FASD United Language Guide link",
        "(Currently absent on the home page.)",
        "Add: \"The words we choose shape whether people feel seen or blamed — and whether they stay "
        "connected to care. Stigmatizing language keeps families from seeking support, reinforces shame, "
        "and obscures the neurological reality of FASD. This guide from FASD United helps you recognize "
        "harmful language and shift toward affirming communication around FASD, substance use, and lived "
        "experience.\" + link to FASD United Language & Stigma Guide.",
        "todo",
    )

    # ========== ABOUT US ==========
    add_heading(doc, "About Us Page (about-us.html)", 2)

    add_heading(doc, "Section Structure", 3)
    add_change_block(
        doc,
        "Page sections / order",
        "Current sections: Who We Are → Mental Health Work Group Members → How to Connect → "
        "Accessibility and Inclusion",
        "Restructure to: Who We Are → Our Approach (NEW) → Our Team (NEW) → Our Partnerships "
        "(REPLACES \"Mental Health Work Group Members\") → How We Work / About This Resource (NEW) → "
        "How to Connect → Accessibility & Inclusion",
        "todo",
        notes="Adds three new sections and renames one.",
    )

    add_heading(doc, "Who We Are — Add Context", 3)
    add_change_block(
        doc,
        "\"Who We Are\" intro (current)",
        "The FASD United Mental Health Work Group is a mental health collaborative dedicated to "
        "advancing brain-based, neurodevelopmentally informed care for individuals with FASD and "
        "other brain-based differences.",
        "ADD before existing text: \"This work is developed in partnership with organizations, "
        "clinicians, researchers, advocates, caregivers, and individuals with lived experience who "
        "believe this care shouldn't be this hard to find and share a commitment to making the shift "
        "to brain-based care more accessible.\" Then keep existing text.",
        "todo",
        notes="Gives first-time visitors context about what the site is and why it matters.",
    )

    add_change_block(
        doc,
        "\"Resources are reviewed\" line",
        "All resources are reviewed by qualified professionals and regularly updated to reflect "
        "current evidence-based practices.",
        "REPLACE with: \"All resources are reviewed by clinicians and researchers with FASD expertise "
        "and regularly updated to reflect current evidence — organized through the Brain-Based Pathway "
        "framework.\" + add a \"last updated [DATE]\" line because the page commits to regular updates.",
        "todo",
    )

    add_heading(doc, "Our Approach — New Section", 3)
    add_addition_block(
        doc,
        "\"Our Approach\" (NEW section)",
        "(Currently absent on About Us page.)",
        "ADD: \"The Brain-Based Pathway is the central framework guiding our professional resources. "
        "It is a developmental roadmap for trauma-informed, neurobehavioral practice, supporting "
        "clinicians from foundational brain-based understanding to advanced, responsive care across "
        "six competency stages. By integrating FASD, neurodevelopmental science, trauma, and "
        "neurodiversity-affirming practice, the pathway promotes individualized, compassionate, and "
        "effective mental health support for people whose brains work differently. This site exists "
        "because we believe the shift to brain-based practice is worth making — and that clinicians, "
        "caregivers, and self-advocates shouldn't have to make it alone.\" + link to Pathway page.",
        "todo",
        notes="Janessa noted the larger workgroup gave thumbs up to this previously.",
    )

    add_heading(doc, "Our Team — New Section", 3)
    add_addition_block(
        doc,
        "\"Our Team\" (NEW section)",
        "(Currently absent — co-leaders not named on the About Us page; only org logos appear under "
        "\"Mental Health Work Group Members\".)",
        "ADD: \"[Co-leaders here: Name, LCPC, Org | 1–2 sentence bio] [Key private practice professionals: "
        "Name, credentials | focus area] This work is also shaped by a broader advisory group of clinicians, "
        "caregivers, and individuals with lived experience. Advisory members contribute to content review, "
        "resource development, and direction-setting for the site.\"",
        "todo",
    )

    add_heading(doc, "How We Work — New Section (Trust-Building)", 3)
    add_addition_block(
        doc,
        "\"How We Work / About This Resource\" (NEW section)",
        "(Currently absent — no transparency statement on conflict of interest, lived-experience role, "
        "or update cadence.)",
        "ADD: \"The Work Group is co-led by [names/titles] and operates under the umbrella of FASD United "
        "Affiliate Network. Content is developed collaboratively and reviewed by clinicians and researchers "
        "with FASD expertise before publication. Conflict of interest: Content decisions are collective; "
        "some group members are also organizational partners — this transparency ensures objectivity. "
        "Lived experience: Integrated through advisory input on content review and direction-setting "
        "(advisory, not decision-making authority). Updates: Reviewed annually or as evidence emerges; "
        "last update [DATE].\"",
        "todo",
    )

    add_heading(doc, "Our Partnerships — Renamed/Reworded", 3)
    add_change_block(
        doc,
        "\"Mental Health Work Group Members\" intro",
        "This work is informed by a collaborative group of organizations, clinicians, caregivers, and "
        "individuals with lived experience.",
        "RENAME section to \"Our Partnerships.\" REPLACE intro with: \"This work is developed in "
        "partnership with organizations, clinicians, caregivers, and individuals with lived experience "
        "who share a commitment to FASD-informed, neurodevelopmentally informed care.\" "
        "Janessa: current is \"confusing on what these organizations actually do.\"",
        "todo",
    )

    add_addition_block(
        doc,
        "Logo permissions",
        "Currently shows org cards with names: FASD United, The Florida Center for Early Childhood, "
        "Papillion Center, FASD Focus, FASCETS, Formed Families Forward, Flow Counseling, Wonderlands, "
        "+ \"Additional members to be confirmed.\"",
        "Then list Organization Logos. CONFIRM logo permissions.",
        "todo",
    )

    add_heading(doc, "How to Connect — Tier the Contacts", 3)
    add_change_block(
        doc,
        "\"How to Connect\" section",
        "Two cards: \"Training or General Questions\" (FASD United phone + contact form) and "
        "\"Website Feedback\" (FASD United contact form).",
        "Tier into three contact types: General FASD Questions → fasdunited@fasdunited.org; "
        "Specific Contributors → direct contact info or through general contact (decide who monitors); "
        "Website Feedback → general contact.",
        "todo",
    )

    # ========== PROFESSIONAL PAGE ==========
    add_heading(doc, "Professional Page (professionals.html)", 2)

    add_heading(doc, "\"When Nothing Works\" — Soften Wording", 3)
    add_change_block(
        doc,
        "Opening framing",
        "Most therapeutic models were built around typical brain functioning. They have real limits "
        "when the brain works differently.",
        "Most therapeutic models were built around typical brain functioning AND MAY NEED ADAPTATION "
        "WHEN BRAIN-BASED DIFFERENCES ARE PRESENT.",
        "todo",
        notes="Tones down the absolute statement.",
    )

    add_change_block(
        doc,
        "CBT bullet",
        "CBT falling flat because cause-and-effect reasoning is impaired",
        "CBT MAY FALL FLAT WHEN abstract reasoning or generalization are affected, REQUIRING ADAPTATION.",
        "todo",
    )

    add_change_block(
        doc,
        "DBT bullet",
        "DBT skills that do not generalize due to executive functioning limitations",
        "DBT skills MAY NOT GENERALIZE EASILY WITHOUT REPETITION AND ENVIRONMENTAL SUPPORT.",
        "todo",
    )

    add_change_block(
        doc,
        "Behavioral plans bullet",
        "Behavioral plans that backfire when consequences feel arbitrary",
        "Behavioral plans MAY BACKFIRE when consequences feel arbitrary, ABSTRACT, OR DELAYED.",
        "todo",
    )

    add_change_block(
        doc,
        "Closing line of \"When Nothing Works\"",
        "The issue is not your competence. The issue is that your training likely did not teach you to "
        "assess and understand the brain before choosing how to treat.",
        "ADD after \"The issue is not your competence\": \"YOUR TRAINING REMAINS VALUABLE. THIS IS ABOUT "
        "ADDING A NEURODEVELOPMENTAL LENS THAT MANY STANDARD MODELS WERE NOT EXPLICITLY DESIGNED TO HIGHLIGHT.\"",
        "todo",
    )

    add_heading(doc, "Step 1 — Soften Graduate Training Statement", 3)
    add_change_block(
        doc,
        "Step 1 opening",
        "Most graduate training in mental health does not include meaningful content on neurodevelopment "
        "or the neurobiology of regulation.",
        "MANY graduate programs PROVIDE LIMITED training in neurodevelopment and the neurobiology of regulation.",
        "todo",
    )

    add_heading(doc, "FAQ — Reframe Prevalence Claim", 3)
    add_change_block(
        doc,
        "FAQ #1 (\"I do not work with FASD specifically...\")",
        "There is a strong chance you are already working with someone whose brain was shaped by prenatal "
        "alcohol exposure without either of you knowing it.",
        "BECAUSE FASD IS OFTEN UNDERDIAGNOSED — ESTIMATES SUGGEST AS MANY AS 1 IN 20 PEOPLE MAY BE AFFECTED — "
        "you are likely already working with clients whose brains MAY HAVE BEEN shaped by prenatal alcohol "
        "exposure THAT HAS NOT BEEN IDENTIFIED.",
        "todo",
        notes="Adds a citable prevalence figure and softens certainty.",
    )

    add_heading(doc, "Brain-Based Practice — Add Nuance & Definitions", 3)
    add_addition_block(
        doc,
        "Constraints disclaimer",
        "(Currently absent — no acknowledgment of structural constraints on clinicians.)",
        "ADD: \"Brain-based planning can be constrained by caseloads, billing, and agency expectations, "
        "and advocacy for structural change is often necessary.\"",
        "todo",
    )

    add_addition_block(
        doc,
        "Broaden FASD framing",
        "Page focuses largely on FASD; doesn't explicitly extend to other conditions in body text.",
        "Explicitly note that brain-based principles apply to autism, ADHD, TBI, complex trauma, and "
        "other neurodevelopmental differences.",
        "partial",
        notes="FAQ #1 already mentions \"autism, ADHD, intellectual disabilities, learning differences, "
              "traumatic brain injury, and complex trauma\" — but Janessa wants this in the main page body, not buried in FAQ.",
    )

    add_addition_block(
        doc,
        "What \"brain-based practice\" actually is",
        "(Currently absent — no defining paragraph on the page itself.)",
        "ADD: \"Brain-based practice is not a separate, manualized treatment, but a clinical orientation "
        "that can be added to existing models like CBT, DBT, TF-CBT, EMDR, PCIT, DDP, CPS, and play therapy. "
        "It assumes collaboration with medical and developmental professionals when needed and emphasizes "
        "that this page is about adapting your practice, not making diagnoses. Brain-based practice asks "
        "for neurodevelopmental humility — recognizing that capacity, not just motivation, shapes what "
        "clients can do — and draws on evidence about FASD, complex trauma, and neurodevelopmental conditions "
        "consistent with trauma-informed and neurodiversity-affirming frameworks. It does not replace the "
        "need for diagnosis when indicated, nor does it override clinical or ethical guidelines; it is a "
        "lens to help clinicians adapt evidence-based practices to clients whose brains work differently.\"",
        "todo",
    )

    add_heading(doc, "Page Structure — New Sections", 3)
    add_addition_block(
        doc,
        "\"Who Is This For?\" section",
        "(Currently absent — clinicians have to infer from \"When Nothing Works\".)",
        "Add explicit section naming: FASD, neurodivergent clients, complex developmental trauma, "
        "attachment disruption, clients labeled \"treatment-resistant.\"",
        "todo",
    )

    add_addition_block(
        doc,
        "\"What Is Brain-Based Care?\" intro section",
        "(Currently absent — page jumps from \"When Nothing Works\" into the 7-step pathway.)",
        "ADD a brief \"What Is Brain-Based Care?\" or \"New to brain-based care?\" section before "
        "asking clinicians to engage with the pathway.",
        "todo",
    )

    add_addition_block(
        doc,
        "Informed vs. Responsive distinction",
        "Currently lives only in the linked Self-Assessment and Competencies pages, not on the "
        "Professionals page itself.",
        "Move the Informed vs. Responsive distinction onto the page itself. \"It's the organizing logic "
        "of the whole pathway and clinicians shouldn't have to click away to find it.\"",
        "todo",
    )

    add_addition_block(
        doc,
        "\"What you will gain\" outcomes section",
        "(Currently absent — clinicians don't see concrete outcomes for committing to the pathway.)",
        "ADD concrete outcomes: clinical confidence, real-time adaptation, reduced frustration, "
        "professional community.",
        "todo",
    )

    add_addition_block(
        doc,
        "Identity-shift framing",
        "After \"The issue is not your competence,\" page moves into the framework intro without "
        "broader identity framing.",
        "Connect brain-based practice to trauma-informed care, cultural humility, and "
        "neurodiversity-affirming practice. Frame it as an identity shift, not just skill acquisition.",
        "todo",
    )

    add_addition_block(
        doc,
        "Step 6 \"Growing Together\" labeling",
        "Step 6: \"Growing Together — Reflective supervision, consultation, and community.\" "
        "Apply pillar says \"Join the consultation and reflection group.\"",
        "Either label it \"coming soon\" or add more content-like language to frame it as a real "
        "community with shared language, cross-disciplinary consultation, and peer support.",
        "todo",
    )

    add_heading(doc, "Subtitle Wording", 3)
    add_change_block(
        doc,
        "Step 4 H3",
        "Planning with the Brain in Mind",
        "Brain-Based Treatment Planning",
        "todo",
    )
    add_change_block(
        doc,
        "Step 6 H3",
        "Growing Together",
        "Reflective Supervision & Community",
        "todo",
    )
    add_change_block(
        doc,
        "Step 3 H3",
        "Seeing the Brain in Your Clients",
        "Seeing Brain First — OR — Looking Through a Brain-Based Lens",
        "todo",
    )

    add_heading(doc, "Step 0 \"Why This Matters\"", 3)
    add_change_block(
        doc,
        "Step 0 framing",
        "Step 0 is a full numbered step in the journey, with Learn / Reflect / Apply / Check In pillars.",
        "\"Why This Matters\" — important, but maybe NOT a full Step 0. Works better as an introductory "
        "section rather than a standalone \"step\" with no explicit skill-building or competency.",
        "todo",
    )

    add_heading(doc, "Steps 1 and 2", 3)
    add_change_block(
        doc,
        "Steps 1 and 2",
        "Step 1: Understanding the Brain. Step 2: Trauma, FASD, and the Brain.",
        "COMBINE Steps 1 and 2 into one: \"Understanding the Brain — including how trauma and prenatal "
        "exposure change it.\"",
        "todo",
        notes="Christie Petrenko also flagged the Step 1/2/3 sequence as unclear and supported reorganizing.",
    )

    add_heading(doc, "Self-Assessment Placement", 3)
    add_addition_block(
        doc,
        "Self-assessment surfacing",
        "Self-assessment links live in: Step 6 Apply pillar, plus the dark CTA section near the bottom "
        "of the page (\"Where Are You on the Journey?\").",
        "Self-assessment seems a bit buried; add a SECOND prompt earlier in the page so clinicians "
        "orient to it faster.",
        "todo",
    )

    add_heading(doc, "Diagnosis Gap", 3)
    add_addition_block(
        doc,
        "Naming the diagnosis gap",
        "FAQ touches on FASD being undiagnosed; main page body does not directly name it as a clinical issue.",
        "Add a sentence or two to explicitly name the diagnosis gap as a clinical issue.",
        "todo",
    )

    add_heading(doc, "FAQ Additions", 3)
    add_addition_block(
        doc,
        "FAQ — practical clinician objections",
        "Current FAQs: relevance, trauma-informed care relationship, modality compatibility, sequence flexibility.",
        "Add FAQs addressing: \"How do I realistically do this in my current setting?\" / "
        "\"My clinic won't support this approach — what do I do?\" / "
        "\"I don't have access to a neuropsych eval\" / "
        "\"I don't have an official diagnosis.\"",
        "todo",
    )

    add_heading(doc, "Call to Action", 3)
    add_change_block(
        doc,
        "Page CTA",
        "Dark CTA section near the bottom: \"Take the Self-Assessment\" + \"View Full Competency Definitions.\"",
        "Need a clearer CTA: \"Start with the self-assessment\" and/or \"Join the community.\"",
        "partial",
        notes="\"Take the Self-Assessment\" already exists but Janessa wants more prominence and a community CTA.",
    )

    add_heading(doc, "Sample 10-Minute Video Topics", 3)
    add_addition_block(
        doc,
        "Step 0 / What Is Brain-Based Responsive videos",
        "(Video content not yet built into the page.)",
        "\"Why This Matters\" and \"What Does It Mean to Be Brain-Based Responsive?\" — distinct, "
        "identity-shift focused, worth their own short videos.",
        "todo",
    )

    add_addition_block(
        doc,
        "Steps 1/2 combined — video topics",
        "(Video content not yet built into the page.)",
        "Topics: screening tools, brain domains, barriers to access. "
        "Resource: Neurodiversity-Affirming Language in Clinical Practice (FASD United).",
        "todo",
    )

    add_addition_block(
        doc,
        "Step 3 video topics",
        "(Video content not yet built into the page.)",
        "Topics: assessments for caregivers who are also struggling. "
        "Resource: Screening Questions You Can Use.",
        "todo",
    )

    add_addition_block(
        doc,
        "Step 4 video topics",
        "(Video content not yet built into the page.)",
        "Topics: realistic goals, caregiver as part of the treatment, documentation & rationale for insurance, "
        "interdisciplinary communication.",
        "todo",
    )

    add_addition_block(
        doc,
        "Step 5 video topics",
        "(Video content not yet built into the page.)",
        "Topics: DDP (Dyadic Developmental Psychotherapy), Collaborative Problem Solving (CPS), Play Therapy, "
        "nervous system: sensory and safety cues.",
        "todo",
    )

    add_addition_block(
        doc,
        "Step 6 video topics",
        "(Video content not yet built into the page.)",
        "Topics (maybe not a video): why it matters for complex clients, managing burnout, "
        "reflective practice / case-consult guide.",
        "todo",
    )

    # ========== DDP ==========
    add_heading(doc, "DDP — Therapy Guide (handout-therapy-guide.html)", 2)

    add_para(
        doc,
        "Note: The current site has already incorporated Janessa's main rewording of the DDP entry. "
        "Her \"Possible additional section: Challenges and Adaptations\" is still pending.",
        italic=True, color=GRAY, size=10,
    )

    add_heading(doc, "What It Asks the Brain to Do", 3)
    add_change_block(
        doc,
        "DDP \"What It Asks the Brain to Do\" (current)",
        "DDP creates felt safety in the therapy room so the child's brain can come out of defensive "
        "survival states (fight, flight, freeze) and move into connection and co-regulation. It is built "
        "around the attitude of PACE: Playfulness, Acceptance, Curiosity, and Empathy. The therapist "
        "uses PACE to help the child experience safety, connection, and being understood, rather than "
        "constant defensiveness. Over time, DDP helps the child's brain build new meanings about themselves "
        "and their experiences, and supports the experience of being truly known and understood by another "
        "person — which helps organize and support a developing brain.",
        "Janessa's proposed wording: same as current text. Her edits to the original draft "
        "(\"defensive survival states,\" co-regulation framing, \"build new meanings\") have already been applied.",
        "done",
    )

    add_heading(doc, "Caregiver Involvement", 3)
    add_change_block(
        doc,
        "DDP \"Caregiver Involvement\" (current)",
        "Very high — caregivers are central to DDP. The child's healing is supported through the "
        "caregiver–child relationship as the therapist supports and coaches both toward safety and connection.",
        "Janessa's proposed wording matches current. \"Caregivers are central to DDP\" framing has been applied.",
        "done",
    )

    add_heading(doc, "Brain-Based Consideration", 3)
    add_change_block(
        doc,
        "DDP Brain-Based Consideration (current)",
        "DDP's relational approach is well-suited for children with brain-based differences who also have "
        "co-occurring attachment disruptions — which is common among children with FASD who have histories "
        "of foster care or adoption. The emphasis on the caregiver's inner state (not just techniques) helps "
        "caregivers manage their own frustration and grief. The verbal and narrative elements may need "
        "adaptation for children with significant cognitive, language, and processing differences.",
        "Janessa added \"language, and processing differences\" to the original \"cognitive differences\" line. "
        "That edit has been applied.",
        "done",
    )

    add_heading(doc, "Challenges and Adaptations — Pending New Section", 3)
    add_addition_block(
        doc,
        "\"Challenges and Adaptations\" subsection (NEW)",
        "(Currently absent — DDP card has What It Asks / Caregiver Involvement / Good Fit / Brain-Based "
        "Consideration, but no Challenges section.)",
        "ADD a new subsection covering: PACE relies on co-regulation; social-cognition challenges can "
        "make synchronized, shared-experience interactions difficult. Intersubjectivity is demanding — "
        "joint attention, theory of mind, and reading social cues can be strained. Narrative work may be "
        "limited by memory, sequencing, and readiness to use emotional language or label feelings. "
        "Reflection may be inconsistent when cause-and-effect reasoning or abstract working memory is impaired. "
        "Pacing may feel too fast; less language, more pauses, and concrete anchors (photos, objects, simple "
        "storylines) help. Caregiver training is key — PACE is a relational stance by the caregiver/clinician, "
        "so the cognitive load shifts off the child and onto the regulated, coached adult.",
        "todo",
        notes="This is the only piece of Janessa's DDP feedback not yet incorporated — but it's substantial.",
    )

    # ========== STATUS SUMMARY ==========
    add_heading(doc, "Status Summary", 2)
    add_para(
        doc,
        "Most of Janessa's feedback is still TO DO. The exceptions are the DDP rewording edits, which "
        "are already on the live site. Her DDP \"Challenges and Adaptations\" subsection, plus all of her "
        "Home Page, About Us, and Professional Page suggestions, are pending decisions on whether to apply.",
        italic=True, color=GRAY, size=10,
    )

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
