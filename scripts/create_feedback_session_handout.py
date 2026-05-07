"""Generate the 'What to Expect in Your Feedback Session' Word handout."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

PROJECT = Path(r"C:\Users\Tamra\Documents\Brain-Based Mental Health Website")
LOGO = PROJECT / "images" / "The Florida Center.png"
OUTPUT = Path(r"C:\Users\Tamra\Documents") / "Handout-Feedback-Session.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x6B)
DARK = RGBColor(0x22, 0x22, 0x22)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Set default body font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Logo, centered ---
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_run = logo_p.add_run()
logo_run.add_picture(str(LOGO), width=Inches(1.8))

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("What to Expect in Your Feedback Session")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = NAVY
tr.font.name = "Calibri"

# --- Subtitle ---
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A roadmap for families receiving assessment results")
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = RGBColor(0x55, 0x5B, 0x66)

doc.add_paragraph()  # spacer

# --- Intro ---
intro = doc.add_paragraph()
ir = intro.add_run(
    "Your feedback session is the beginning of the journey, not the end. "
    "This handout walks you through who will be there, what each part of the "
    "session looks like, and what happens after. Please feel free to reach "
    "out at any time with questions."
)
ir.font.size = Pt(11)

doc.add_paragraph()  # spacer

# --- Sections ---
sections = [
    ("Who Joins the Session",
     "The caregiver is the primary person who attends the feedback session. "
     "If there is a situation where a supportive person also needs to hear "
     "the information, they are welcome to join."),
    ("Welcome and Check-In",
     "The clinic lead opens the session, welcomes your family, and checks in "
     "with you. They will reiterate that this is the beginning of the journey, "
     "not the end, and that you should feel free to reach out at any time."),
    ("Introductions",
     "Introductions are only done when there is a new person in the group. "
     "If everyone has met before, we move directly into the content of the session."),
    ("Diagnosis and Mental Health Recommendations",
     "The clinic lead begins by going over the diagnosis and any mental health "
     "related recommendations. This sets the foundation for the rest of the session."),
    ("How You Learn Best",
     "If we have not already asked how you learn best, we will ask now. This "
     "helps us make sure the resources we share match the way your brain "
     "processes information."),
    ("Each Specialist Shares Their Findings",
     None),  # Handled with bullets below
    ("Your Questions",
     "After each specialist shares, there is time set aside for any questions "
     "you have. No question is too small."),
    ("Next Steps",
     None),  # Handled with subheadings below
]

for idx, (heading, body) in enumerate(sections, start=1):
    h = doc.add_paragraph()
    hr = h.add_run(f"{idx}. {heading}")
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)

    if body:
        p = doc.add_paragraph()
        pr = p.add_run(body)
        pr.font.size = Pt(11)
        pr.font.color.rgb = DARK

    # Helpers
    def add_subheading(text):
        sh = doc.add_paragraph()
        sr = sh.add_run(text)
        sr.bold = True
        sr.font.size = Pt(11.5)
        sr.font.color.rgb = NAVY
        sh.paragraph_format.space_before = Pt(6)
        sh.paragraph_format.space_after = Pt(2)

    def add_body(text):
        bp = doc.add_paragraph()
        br = bp.add_run(text)
        br.font.size = Pt(11)
        br.font.color.rgb = DARK

    def add_bullet(text, bold_lead=None):
        b = doc.add_paragraph(style="List Bullet")
        if bold_lead:
            r1 = b.add_run(bold_lead)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = b.add_run(text)
            r2.font.size = Pt(11)
        else:
            r = b.add_run(text)
            r.font.size = Pt(11)

    # Custom bullet sections
    if heading == "Each Specialist Shares Their Findings":
        lead = doc.add_paragraph()
        lr = lead.add_run(
            "Each specialist will give an overview of the testing they "
            "completed and walk you through their feedback and resources. "
            "They share in this order:"
        )
        lr.font.size = Pt(11)
        for label, desc in [
            ("Neuropsychology", "testing, feedback, and resources."),
            ("Speech-Language Pathology (SLP)", "same format, covering testing, feedback, and resources."),
            ("Occupational Therapy (OT)", "same format, covering testing, feedback, and resources."),
        ]:
            b = doc.add_paragraph(style="List Bullet")
            r1 = b.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = b.add_run(desc)
            r2.font.size = Pt(11)

    if heading == "Next Steps":
        add_body(
            "Next steps will be reviewed either at the end of the feedback "
            "session or directly with the clinic lead. You will leave with a "
            "clear sense of what comes next. Here are some of the things to "
            "focus on:"
        )

        add_subheading("Your Report Timeline")
        add_body(
            "The clinic lead will explain when to expect the preliminary "
            "report and the final report, so you know what is coming and when."
        )

        add_subheading("Understanding Your Report")
        add_body(
            "There will be a lot of information to go through, but when you "
            "get the final report, there are a couple of things as first "
            "steps (depending on what is going on):"
        )
        add_bullet(
            "this will focus on getting the diagnosis medically confirmed "
            "and a summary of any possible referrals we recommend in this "
            "testing. (this may not apply to all families)",
            bold_lead="What to do with this report with your medical professional: ",
        )
        add_bullet(
            "",
            bold_lead="What to do with this report with the school.",
        )

        add_subheading("Follow-Up from Crystal")
        add_body(
            "Crystal will follow up with your family twice: once at two weeks "
            "and again at two months after the feedback session. If you hear "
            "from her, please do not ignore her. These check-ins are part of "
            "how we stay connected to your family through the next stage of "
            "the journey. She also can support if needed with the educational "
            "piece."
        )

        add_subheading("Satisfaction Survey")
        add_body(
            "A satisfaction survey will be sent to you through Foxit. The "
            "clinic lead will explain this at the end of the session. Your "
            "feedback helps us continue to improve the experience for other "
            "families."
        )

# --- Closing reminder ---
doc.add_paragraph()
closer = doc.add_paragraph()
closer.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = closer.add_run(
    "This feedback session is the start of the path forward, not the finish "
    "line. Reach out whenever you need to. We are here for the long journey, "
    "not just today."
)
cr.italic = True
cr.font.size = Pt(11)
cr.font.color.rgb = NAVY

doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
