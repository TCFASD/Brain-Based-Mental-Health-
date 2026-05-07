"""Convert Team-Feedback-Summary-By-Section.md to a Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "Group Feedback" / "Team-Feedback-Summary-By-Section.md"
OUT = ROOT / "Group Feedback" / "Team-Feedback-Summary-By-Section.docx"

GREEN = RGBColor(0x4A, 0x6B, 0x5C)
DARK = RGBColor(0x2E, 0x2E, 0x2E)


def add_runs_with_inline(paragraph, text):
    """Render **bold**, *italic*, `code`, and [reviewer] inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\])")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
        elif token.startswith("[") and token.endswith("]"):
            r = paragraph.add_run(token)
            r.italic = True
            r.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def style_heading(p, level):
    for run in p.runs:
        run.font.color.rgb = GREEN
        run.bold = True
    if level == 1:
        for run in p.runs:
            run.font.size = Pt(20)
    elif level == 2:
        for run in p.runs:
            run.font.size = Pt(15)
    elif level == 3:
        for run in p.runs:
            run.font.size = Pt(12)


def main():
    md = MD.read_text(encoding="utf-8")
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    lines = md.split("\n")
    i = 0
    in_table = False
    table_rows = []
    while i < len(lines):
        line = lines[i].rstrip()

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.add_run(line[2:].strip())
            style_heading(p, 1)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.add_run(line[3:].strip())
            style_heading(p, 2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.add_run(line[4:].strip())
            style_heading(p, 3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Tables (start with `| ... |` and next line `|---|---|`)
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|\s*$", lines[i + 1]):
            # Parse header
            headers = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                rows.append(row)
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
            for r, row in enumerate(rows, start=1):
                for c, val in enumerate(row):
                    if c < len(headers):
                        cell = table.rows[r].cells[c]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs_with_inline(p, val)
            doc.add_paragraph()
            continue

        # Bulleted lists
        if line.lstrip().startswith("- "):
            indent_chars = len(line) - len(line.lstrip())
            level = indent_chars // 2
            content = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
            add_runs_with_inline(p, content)
            i += 1
            continue

        # Numbered lists
        m_num = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline(p, m_num.group(2))
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_runs_with_inline(p, line)
        i += 1

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
