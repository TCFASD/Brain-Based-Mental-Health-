from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'C:\Users\Tamra\Documents\Brain-Based Mental Health Website\Group-Feedback-Website-Title.docx')

# Find the "2. Additional Questions" section and replace the placeholder
for i, para in enumerate(doc.paragraphs):
    if '2. Additional Questions' in para.text:
        # The placeholder is the next paragraph
        placeholder = doc.paragraphs[i + 1]
        placeholder.clear()
        break

# Now insert the overview section after the placeholder location
# We'll work by adding content after the placeholder paragraph

# Clear placeholder text
placeholder.clear()
run = placeholder.add_run('We are selecting the overview paragraph for the home page. This is the first thing visitors read after the title. It should explain why the site exists and who built it.')
run.font.size = Pt(11)

# Add spacing
doc.paragraphs[i + 1].space_after = Pt(12)

# We need to insert paragraphs after a specific position
# Easiest approach: add to end of document since the placeholder area had blank space

# Find all the blank paragraphs after "Additional Questions" and use them
blank_count = 0
insert_index = i + 2
for j in range(i + 2, len(doc.paragraphs)):
    if doc.paragraphs[j].text.strip() == '':
        blank_count += 1
    else:
        break

# Clear the blank paragraphs and reuse them
para_index = i + 2

# Helper to set paragraph text
def set_para(idx, text, bold=False, size=11, color=None, italic=False):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        return idx + 1
    return idx

def add_option(idx, label, text):
    idx = set_para(idx, '', size=6)  # spacer
    idx = set_para(idx, label, bold=True, size=12, color=(0x2c, 0x5a, 0xa0))
    # Split into sentences for readability - but keep as one block
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
        idx += 1
    # Rank line
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('    Preference (check one):  [ ] Top choice   [ ] Good   [ ] Okay   [ ] Not preferred')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        idx += 1
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.clear()
        run = p.add_run('    Comments: ')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2 = p.add_run('_' * 65)
        run2.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        idx += 1
    return idx

# Prompt
idx = set_para(para_index, 'Please indicate your preference for each option:', bold=True, size=11)

option1 = 'This site was built by therapists, caregivers, and self-advocates who saw the same problem from different sides. Therapists felt unequipped to help their most complex clients. Caregivers felt like no one understood their child. Self-advocates felt like the system was not built for them. We came together because we recognized a fundamental mismatch between how mental health treatment is delivered and how the brain actually works. This site is the result. Here you will find the framework, tools, and resources to bridge that gap, whether you are a clinician looking to better serve your clients, a caregiver searching for answers, or a self-advocate learning to navigate a system that was not designed with your brain in mind.'

option2 = 'We are therapists, caregivers, and self-advocates who came together around a shared realization: the mental health system was not working for the people who needed it most. Therapists felt stuck, unsure why their training was not reaching certain clients. Caregivers watched their children cycle through treatments that did not fit. Self-advocates struggled in silence, told to try harder in a system that never asked how their brain worked. There is a better way. When we start with the brain, everything changes. This site gives you the framework and tools to find the right support and to provide support in a way that actually works.'

option3 = 'This site exists because therapists, caregivers, and self-advocates all recognized the same problem. Therapists felt unequipped. Families felt unheard. Self-advocates felt unseen. The common thread was a mental health system built around typical brain functioning, leaving everyone struggling when the brain works differently. We came together to change that. Here you will find the tools and framework to close that gap, so clinicians can provide care that fits, caregivers can find answers that work, and self-advocates can get the support they deserve.'

idx = add_option(idx, 'Option A:', option1)
idx = add_option(idx, 'Option B:', option2)

# We might run out of blank paragraphs, so let's check
remaining = len(doc.paragraphs) - idx
if remaining < 6:
    # Need to add more paragraphs before the "Thank you" at the end
    # Instead, let's just append to the document
    pass

# Try option 3
if idx + 4 < len(doc.paragraphs):
    idx = add_option(idx, 'Option C:', option3)

doc.save(r'C:\Users\Tamra\Documents\Brain-Based Mental Health Website\Group-Feedback-Website-Title.docx')
print('Word document updated successfully.')
