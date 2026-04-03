from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)

def heading(text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
        p.paragraph_format.space_before = Pt(24)
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
        p.paragraph_format.space_before = Pt(12)
    return p

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def option_block(label, text):
    doc.add_paragraph()
    heading(label, level=3)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run('    Preference:  [ ] Top choice   [ ] Good   [ ] Okay   [ ] Not preferred')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p = doc.add_paragraph()
    run = p.add_run('    Comments: ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2 = p.add_run('_' * 65)
    run2.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

def blank_lines(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        run = p.add_run('')
        run.font.size = Pt(6)

def response_lines(n=3):
    for _ in range(n):
        p = doc.add_paragraph()
        run = p.add_run('_' * 85)
        run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        run.font.size = Pt(10)

# ============================================================
# DOCUMENT START
# ============================================================

heading('Brain-Based Mental Health Community')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Website Development \u2014 Group Feedback')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Please review and share your thoughts on the following questions.\nYour input will shape the direction of our combined website.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

blank_lines(1)

# ============================================================
# SECTION 1: TITLE
# ============================================================
heading('1. Website Title', level=2)

body('We are selecting a title for the combined website. The title should communicate that this resource is for everyone \u2014 clinicians, caregivers, and self-advocates \u2014 and that understanding the brain is the starting point for effective mental health support.')

body('Please rank your preference (1 = top choice, 4 = least preferred) or note any reactions:')

option_block('Option A',
    'Brain-Based Mental Health: Better Outcomes Start with Understanding the Brain')

option_block('Option B',
    'When We Start with the Brain, Everyone Benefits')

option_block('Option C',
    'The Brain-Based Approach: Connecting Clinicians, Caregivers, and Self-Advocates')

option_block('Option D',
    'It Starts with the Brain: Mental Health That Actually Fits')

blank_lines(1)
body('Other title suggestion:')
response_lines(2)

blank_lines(2)

# ============================================================
# SECTION 2: OVERVIEW PARAGRAPH
# ============================================================
heading('2. Home Page Overview Paragraph', level=2)

body('The overview paragraph is the first thing visitors read after the title. It explains why the site exists and who built it. Please rank your preference:')

option_block('Option A',
    'This site was built by therapists, caregivers, and self-advocates who saw the same problem from different sides. Therapists felt unequipped to help their most complex clients. Caregivers felt like no one understood their child. Self-advocates felt like the system was not built for them. We came together because we recognized a fundamental mismatch between how mental health treatment is delivered and how the brain actually works. This site is the result. Here you will find the framework, tools, and resources to bridge that gap, whether you are a clinician looking to better serve your clients, a caregiver searching for answers, or a self-advocate learning to navigate a system that was not designed with your brain in mind.')

option_block('Option B',
    'We are therapists, caregivers, and self-advocates who came together around a shared realization: the mental health system was not working for the people who needed it most. Therapists felt stuck, unsure why their training was not reaching certain clients. Caregivers watched their children cycle through treatments that did not fit. Self-advocates struggled in silence, told to try harder in a system that never asked how their brain worked. There is a better way. When we start with the brain, everything changes. This site gives you the framework and tools to find the right support and to provide support in a way that actually works.')

option_block('Option C',
    'This site exists because therapists, caregivers, and self-advocates all recognized the same problem. Therapists felt unequipped. Families felt unheard. Self-advocates felt unseen. The common thread was a mental health system built around typical brain functioning, leaving everyone struggling when the brain works differently. We came together to change that. Here you will find the tools and framework to close that gap, so clinicians can provide care that fits, caregivers can find answers that work, and self-advocates can get the support they deserve.')

blank_lines(1)
body('Comments or suggested edits to any of the options above:')
response_lines(3)

blank_lines(2)

# ============================================================
# SECTION 3: ABOUT US - HOW TO CONNECT
# ============================================================
heading('3. About Us Page \u2014 How to Connect', level=2)

body('We need to determine who should be listed as the contact for different types of inquiries. Currently the working draft lists FASD United as the placeholder.')

heading('Question 3a: Training Inquiries', level=3)
body('Who should be the point of contact for training inquiries? Should this be FASD United, The Florida Center, or someone else?')
response_lines(2)

heading('Question 3b: General Questions', level=3)
body('Who should be the point of contact for general questions about the site and its content?')
response_lines(2)

heading('Question 3c: Specific Contributors', level=3)
body('Should individual contributors be listed with contact information? If so, who?')
response_lines(2)

heading('Question 3d: Website Feedback', level=3)
body('Who should receive website feedback and suggestions for improvement?')
response_lines(2)

blank_lines(2)

# ============================================================
# SECTION 4: ABOUT US - ORGANIZATIONS IN THE WORKGROUP
# ============================================================
heading('4. About Us Page \u2014 Mental Health Work Group Members', level=2)

body('We want to list all organizations and members of the Mental Health Work Group on the About Us page. Below is our current list. Please confirm, correct, or add to it.')

heading('Organizations (current list \u2014 please add any that are missing):', level=3)

orgs = [
    'FASD United',
    'The Florida Center for Early Childhood',
    'Papillion Center',
    'FASD Focus',
    'FASCETS',
    'Formed Families Forward',
    'Flow Counseling',
    'Wonderlands',
]

for org in orgs:
    p = doc.add_paragraph()
    run = p.add_run('    [ ] Confirm   ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run = p.add_run(org)
    run.font.size = Pt(11)

blank_lines(1)
body('Additional organizations to add:')
response_lines(3)

blank_lines(1)
heading('Question 4a: Private Practice Clinicians', level=3)
body('How should we list private practice clinicians who are part of the group? By name and credentials? By practice name? Both? Should we create a separate section for individual clinicians versus organizations?')
response_lines(3)

heading('Question 4b: Caregivers in the Work Group', level=3)
body('How should caregivers who are part of the work group be represented on the About Us page? By name? As a group acknowledgment? Other?')
response_lines(3)

blank_lines(2)

# ============================================================
# SECTION 5: OUR APPROACH PLACEMENT
# ============================================================
heading('5. "Our Approach" Section Placement', level=2)

body('The current "Our Approach" text describes the Brain-Based Pathway framework. It focuses on clinicians and the six competency stages. We are considering where this section belongs on the website.')

body('The current text reads:')

body('"The Brain-Based Pathway is the central framework guiding our professional resources. It is a developmental roadmap for trauma-informed, neurobehavioral practice, supporting clinicians from foundational brain-based understanding to advanced, responsive care across six competency stages. By integrating FASD, neurodevelopmental science, trauma, and neurodiversity-affirming practice, the pathway promotes individualized, compassionate, and effective mental health support for people whose brains work differently. This site exists because we believe the shift to brain-based practice is worth making and that clinicians, caregivers, and self-advocates should not have to make it alone."', italic=True)

heading('Question 5a:', level=3)
body('Since this text is focused on clinicians, should "Our Approach" live on the Professionals page instead of the About Us page? Or should we rewrite it so it speaks to all three audiences (clinicians, caregivers, and self-advocates)?')

p = doc.add_paragraph()
run = p.add_run('    [ ] Move to Professionals page as-is')
run.font.size = Pt(11)
p = doc.add_paragraph()
run = p.add_run('    [ ] Keep on About Us but rewrite to include all three audiences')
run.font.size = Pt(11)
p = doc.add_paragraph()
run = p.add_run('    [ ] Other (explain below)')
run.font.size = Pt(11)

blank_lines(1)
body('Comments:')
response_lines(3)

blank_lines(2)

# ============================================================
# SECTION 6: ADDITIONAL QUESTIONS (PLACEHOLDER)
# ============================================================
heading('6. Additional Questions', level=2)

body('(Space reserved for additional questions to be added)', italic=True, color=(0x99, 0x99, 0x99))
blank_lines(6)

# ============================================================
# FOOTER
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Thank you for your feedback!')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Please return your responses to: ___________________')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(r'C:\Users\Tamra\Documents\Brain-Based Mental Health Website\Group-Feedback-Website-Title.docx')
print('Word document rebuilt successfully.')
